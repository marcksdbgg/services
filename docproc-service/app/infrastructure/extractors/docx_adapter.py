import io
import docx  # python-docx
import structlog
from typing import Tuple, Dict, Any, Union, List

from app.application.ports.extraction_port import ExtractionPort, ExtractionError, UnsupportedContentTypeError
from app.infrastructure.extractors.base_extractor import BaseExtractorAdapter

log = structlog.get_logger(__name__)

class DocxAdapter(BaseExtractorAdapter):
    """Adaptador para extraer texto de archivos DOCX."""

    SUPPORTED_CONTENT_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword" # python-docx can sometimes handle .doc, but it's not guaranteed
    ]

    def extract_text(
        self,
        file_bytes: bytes,
        filename: str,
        content_type: str
    ) -> Tuple[str, Dict[str, Any]]:
        if content_type not in self.SUPPORTED_CONTENT_TYPES:
            raise UnsupportedContentTypeError(f"DocxAdapter does not support content type: {content_type}")

        log.debug("DocxAdapter: Extracting text from DOCX bytes", filename=filename)
        extraction_metadata: Dict[str, Any] = {}
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs if p.text and not p.text.isspace()])
            
            extraction_metadata["num_paragraphs_extracted"] = len([p for p in doc.paragraphs if p.text and not p.text.isspace()])
            log.info("DocxAdapter: DOCX extraction successful", filename=filename, num_paragraphs=extraction_metadata["num_paragraphs_extracted"])
            return text, extraction_metadata
        except Exception as e:
            # If it's a .doc file, it might fail here. Log a specific warning.
            if content_type == "application/msword":
                log.warning("DocxAdapter: Failed to process .doc file. This format has limited support.", filename=filename, error=str(e))
            raise self._handle_extraction_error(e, filename, "DocxAdapter")